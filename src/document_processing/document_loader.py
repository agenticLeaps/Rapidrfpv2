import os
import logging
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import PyPDF2
from docx import Document
import tiktoken

logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any]
    token_count: int

@dataclass
class ProcessedDocument:
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    total_tokens: int

class DocumentLoader:
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.tokenizer = tiktoken.get_encoding("cl100k_base")  # GPT-4 tokenizer
    
    def load_document(self, file_path: str) -> Optional[ProcessedDocument]:
        """Load and process a document from file path."""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                text = self._load_pdf(file_path)
            elif file_extension == '.docx':
                text = self._load_docx(file_path)
            elif file_extension in ['.txt', '.md']:
                text = self._load_text(file_path)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
            
            if not text or not text.strip():
                logger.warning(f"No text content found in {file_path}")
                return None
            
            # Create document metadata
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'file_extension': file_extension,
                'file_size': os.path.getsize(file_path),
                'total_chars': len(text)
            }
            
            # Create chunks
            chunks = self._create_chunks(text, metadata)
            total_tokens = sum(chunk.token_count for chunk in chunks)
            
            logger.info(f"Loaded document {file_path}: {len(chunks)} chunks, {total_tokens} tokens")
            
            return ProcessedDocument(
                chunks=chunks,
                metadata=metadata,
                total_tokens=total_tokens
            )
            
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            return None
    
    def _load_pdf(self, file_path: str) -> str:
        """Load text from PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n[Page {page_num + 1}]\n{page_text}"
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return ""
    
    def _load_docx(self, file_path: str) -> str:
        """Load text from DOCX file."""
        try:
            doc = Document(file_path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return ""
    
    def _load_text(self, file_path: str) -> str:
        """Load text from TXT or MD file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                return ""
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""
    
    def _create_chunks(self, text: str, doc_metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split text into overlapping chunks based on token count."""
        # Tokenize the entire text
        tokens = self.tokenizer.encode(text)
        
        if len(tokens) <= self.chunk_size:
            # Document is small enough to be a single chunk
            return [DocumentChunk(
                content=text,
                chunk_index=0,
                start_char=0,
                end_char=len(text),
                metadata={**doc_metadata, 'chunk_index': 0, 'total_chunks': 1},
                token_count=len(tokens)
            )]
        
        chunks = []
        chunk_index = 0
        start_token = 0
        
        while start_token < len(tokens):
            # Calculate end token for this chunk
            end_token = min(start_token + self.chunk_size, len(tokens))
            
            # Extract chunk tokens
            chunk_tokens = tokens[start_token:end_token]
            
            # Decode back to text
            chunk_text = self.tokenizer.decode(chunk_tokens)
            
            # Find character positions (approximate)
            start_char = self._estimate_char_position(text, start_token, tokens)
            end_char = self._estimate_char_position(text, end_token, tokens)
            
            # Create chunk metadata
            chunk_metadata = {
                **doc_metadata,
                'chunk_index': chunk_index,
                'start_token': start_token,
                'end_token': end_token,
                'token_count': len(chunk_tokens)
            }
            
            chunks.append(DocumentChunk(
                content=chunk_text.strip(),
                chunk_index=chunk_index,
                start_char=start_char,
                end_char=end_char,
                metadata=chunk_metadata,
                token_count=len(chunk_tokens)
            ))
            
            # Move to next chunk with overlap
            if end_token >= len(tokens):
                break
            
            start_token = end_token - self.chunk_overlap
            chunk_index += 1
        
        # Update total chunks in metadata
        for chunk in chunks:
            chunk.metadata['total_chunks'] = len(chunks)
        
        return chunks
    
    def _estimate_char_position(self, text: str, token_position: int, all_tokens: List[int]) -> int:
        """Estimate character position from token position."""
        if token_position == 0:
            return 0
        if token_position >= len(all_tokens):
            return len(text)
        
        # Decode up to the token position to estimate character position
        partial_tokens = all_tokens[:token_position]
        partial_text = self.tokenizer.decode(partial_tokens)
        return len(partial_text)
    
    def get_chunk_context(self, chunk: DocumentChunk, context_chunks: int = 2) -> str:
        """Get surrounding context for a chunk (not implemented in this basic version)."""
        # This would require access to the full document chunks
        # For now, just return the chunk content
        return chunk.content
    
    def estimate_processing_cost(self, file_path: str) -> Dict[str, Any]:
        """Estimate the processing cost for a document."""
        if not os.path.exists(file_path):
            return {'error': 'File not found'}
        
        file_size = os.path.getsize(file_path)
        
        # Rough estimation based on file size
        estimated_chars = file_size  # Rough estimate
        estimated_tokens = estimated_chars // 4  # Rough tokens per character
        estimated_chunks = max(1, estimated_tokens // self.chunk_size)
        
        # Estimate LLM calls needed
        extraction_calls = estimated_chunks  # One call per chunk for extraction
        embedding_calls = estimated_chunks * 3  # Rough estimate for S, A, H nodes
        
        return {
            'estimated_chars': estimated_chars,
            'estimated_tokens': estimated_tokens,
            'estimated_chunks': estimated_chunks,
            'estimated_llm_calls': extraction_calls,
            'estimated_embedding_calls': embedding_calls,
            'file_size_mb': file_size / (1024 * 1024)
        }